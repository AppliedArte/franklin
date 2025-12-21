"""Document ingestion pipeline - parse, chunk, embed, store."""

import uuid
from io import BytesIO
from typing import Optional

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from src.db.models import Document, DocumentChunk, UserFact
from src.rag.embeddings import EmbeddingService


class TextChunker:
    """Split text into overlapping chunks for embedding."""

    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def chunk_text(self, text: str) -> list[dict]:
        """Split text into chunks with metadata."""
        if not text.strip():
            return []

        chunks = []
        start = 0
        chunk_index = 0

        while start < len(text):
            end = start + self.chunk_size

            # Try to break at sentence/paragraph boundary
            if end < len(text):
                # Look for paragraph break first
                para_break = text.rfind("\n\n", start, end)
                if para_break > start + self.chunk_size // 2:
                    end = para_break + 2
                else:
                    # Look for sentence break
                    for sep in [". ", ".\n", "! ", "? "]:
                        sent_break = text.rfind(sep, start, end)
                        if sent_break > start + self.chunk_size // 2:
                            end = sent_break + len(sep)
                            break

            chunk_text = text[start:end].strip()
            if chunk_text:
                chunks.append({
                    "content": chunk_text,
                    "chunk_index": chunk_index,
                    "start_char": start,
                    "end_char": end,
                })
                chunk_index += 1

            # Move start position with overlap
            start = end - self.chunk_overlap
            if start >= len(text):
                break

        return chunks


class DocumentParser:
    """Parse different document formats to text."""

    @staticmethod
    async def parse_pdf(file_bytes: bytes) -> str:
        """Parse PDF to text using pypdf."""
        try:
            from pypdf import PdfReader

            reader = PdfReader(BytesIO(file_bytes))
            text_parts = []
            for page_num, page in enumerate(reader.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"[Page {page_num}]\n{page_text}")
            return "\n\n".join(text_parts)
        except ImportError:
            raise ImportError("pypdf is required for PDF parsing. Install with: pip install pypdf")

    @staticmethod
    async def parse_docx(file_bytes: bytes) -> str:
        """Parse DOCX to text."""
        try:
            import docx

            doc = docx.Document(BytesIO(file_bytes))
            return "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        except ImportError:
            raise ImportError("python-docx is required for DOCX parsing. Install with: pip install python-docx")

    @staticmethod
    async def parse_txt(file_bytes: bytes) -> str:
        """Parse plain text."""
        return file_bytes.decode("utf-8", errors="ignore")

    async def parse(self, file_bytes: bytes, mime_type: str) -> str:
        """Parse document based on MIME type."""
        if "pdf" in mime_type:
            return await self.parse_pdf(file_bytes)
        elif "word" in mime_type or "docx" in mime_type:
            return await self.parse_docx(file_bytes)
        elif "text" in mime_type:
            return await self.parse_txt(file_bytes)
        else:
            # Try as plain text fallback
            return await self.parse_txt(file_bytes)


class DocumentIngestionService:
    """Full document ingestion pipeline."""

    def __init__(self):
        self.parser = DocumentParser()
        self.chunker = TextChunker()
        self.embedder = EmbeddingService()

    async def ingest_document(
        self,
        db: AsyncSession,
        user_id: str,
        file_bytes: bytes,
        filename: str,
        mime_type: str,
        doc_type: str = "other",
        title: Optional[str] = None,
        description: Optional[str] = None,
    ) -> Document:
        """Ingest a document: parse, chunk, embed, store."""
        # Create document record
        doc = Document(
            id=str(uuid.uuid4()),
            user_id=user_id,
            filename=filename,
            mime_type=mime_type,
            doc_type=doc_type,
            title=title or filename,
            description=description,
            file_size_bytes=len(file_bytes),
            is_processed=False,
        )
        db.add(doc)
        await db.flush()

        try:
            # Parse document to text
            text = await self.parser.parse(file_bytes, mime_type)

            if not text.strip():
                doc.processing_error = "No text content extracted"
                doc.is_processed = True
                return doc

            # Chunk the text
            chunks = self.chunker.chunk_text(text)

            if not chunks:
                doc.processing_error = "No chunks generated"
                doc.is_processed = True
                return doc

            # Generate embeddings for all chunks
            chunk_texts = [c["content"] for c in chunks]
            embeddings = await self.embedder.embed_texts_batch(chunk_texts)

            # Create chunk records
            for chunk_data, embedding in zip(chunks, embeddings):
                chunk = DocumentChunk(
                    id=str(uuid.uuid4()),
                    document_id=doc.id,
                    content=chunk_data["content"],
                    chunk_index=chunk_data["chunk_index"],
                    embedding=embedding,
                )
                db.add(chunk)

            doc.chunk_count = len(chunks)
            doc.is_processed = True

        except Exception as e:
            doc.processing_error = str(e)
            doc.is_processed = True

        return doc

    async def ingest_user_fact(
        self,
        db: AsyncSession,
        user_id: str,
        category: str,
        fact: str,
        source: str = "conversation",
        conversation_id: Optional[str] = None,
        confidence: float = 1.0,
    ) -> UserFact:
        """Store a fact learned about the user with embedding."""
        # Generate embedding for the fact
        embedding = await self.embedder.embed_text(fact)

        user_fact = UserFact(
            id=str(uuid.uuid4()),
            user_id=user_id,
            category=category,
            fact=fact,
            source=source,
            conversation_id=conversation_id,
            confidence=confidence,
            embedding=embedding,
        )
        db.add(user_fact)

        return user_fact

    async def ingest_text_content(
        self,
        db: AsyncSession,
        user_id: str,
        content: str,
        filename: str = "text_input",
        doc_type: str = "other",
        title: Optional[str] = None,
    ) -> Document:
        """Ingest plain text content (e.g., pasted text, notes)."""
        return await self.ingest_document(
            db=db,
            user_id=user_id,
            file_bytes=content.encode("utf-8"),
            filename=filename,
            mime_type="text/plain",
            doc_type=doc_type,
            title=title,
        )

    async def close(self):
        """Cleanup resources."""
        await self.embedder.close()
